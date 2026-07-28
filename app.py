from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import tempfile
import PyPDF2
from docx import Document
from openpyxl import Workbook, load_workbook
import pdf2docx
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import traceback
import re
import io
import base64

# ============ OCR LIBRARY ============
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
    print("✅ OCR Library loaded successfully")
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"⚠️ OCR Library not available: {e}")

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

# ============ OCR FUNCTION ============
def extract_text_with_ocr(pdf_path):
    """Ekstrak teks dari PDF scan/foto menggunakan OCR"""
    if not OCR_AVAILABLE:
        return None
    
    try:
        print("🔍 Starting OCR process...")
        
        # Convert PDF ke gambar dengan resolusi tinggi
        images = pdf2image.convert_from_path(
            pdf_path, 
            dpi=300,
            first_page=1,
            last_page=10  # Batasi 10 halaman pertama untuk performa
        )
        
        all_text = ""
        for i, image in enumerate(images):
            print(f"📄 Processing page {i+1}/{len(images)}")
            
            # Preprocessing untuk hasil OCR lebih baik
            gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
            
            # Thresholding (menjadi hitam-putih)
            _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # OCR dengan bahasa Indonesia + Inggris
            custom_config = r'--oem 3 --psm 6 -l ind+eng'
            text = pytesseract.image_to_string(thresh, config=custom_config)
            
            if text.strip():
                all_text += f"Halaman {i+1}:\n{text}\n\n"
            else:
                # Coba lagi tanpa preprocessing
                text2 = pytesseract.image_to_string(image, lang='ind+eng')
                if text2.strip():
                    all_text += f"Halaman {i+1}:\n{text2}\n\n"
        
        if not all_text.strip():
            print("⚠️ No text extracted by OCR")
            return None
            
        print(f"✅ OCR complete: {len(all_text)} characters extracted")
        return all_text
        
    except Exception as e:
        print(f"❌ OCR Error: {str(e)}")
        print(traceback.format_exc())
        return None

# ============ DETEKSI TABEL DARI TEKS ============
def detect_table_from_text(text):
    """Deteksi dan ekstrak tabel dari teks"""
    lines = text.split('\n')
    table_data = []
    current_row = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Coba deteksi kolom dengan spasi banyak atau tab
        parts = re.split(r'\s{3,}|\t', line)
        
        if len(parts) >= 3:
            # Ini kemungkinan baris tabel
            row = [p.strip() for p in parts if p.strip()]
            table_data.append(row)
        elif current_row:
            # Gabung dengan baris sebelumnya
            current_row.append(line)
            if len(current_row) >= 3:
                table_data.append(current_row)
                current_row = []
        else:
            current_row.append(line)
    
    return table_data

# ============ PDF TO WORD ============
@app.route('/convert/pdf-to-word', methods=['POST'])
def pdf_to_word():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        print(f"📁 Processing file: {file.filename}")
        
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_input.name)
        temp_input.close()
        
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_output.close()
        
        try:
            # STEP 1: Coba ekstrak teks biasa (PDF teks)
            print("📖 Trying regular text extraction...")
            pdf_reader = PyPDF2.PdfReader(temp_input.name)
            text = ''
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + '\n\n'
            
            # STEP 2: Kalau kosong, pake OCR
            if not text.strip():
                print("⚠️ No text found, trying OCR...")
                text = extract_text_with_ocr(temp_input.name)
                
                if text is None or not text.strip():
                    return jsonify({
                        'error': 'Tidak bisa membaca teks dari PDF. Pastikan file tidak rusak atau terlalu buram.',
                        'hint': 'Coba file PDF yang jelas atau hasil scan dengan resolusi tinggi.'
                    }), 400
            
            # STEP 3: Buat Word document
            print("📝 Creating Word document...")
            doc = Document()
            
            # Bersihkan teks
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Hapus karakter aneh
                    line = re.sub(r'[^\x00-\x7F]+', ' ', line)
                    if len(line) > 2:
                        doc.add_paragraph(line)
            
            doc.save(temp_output.name)
            
            print("✅ Word document created successfully")
            
            return send_file(
                temp_output.name,
                as_attachment=True,
                download_name='output.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            print(f"❌ Error in conversion: {str(e)}")
            print(traceback.format_exc())
            return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500
        finally:
            os.unlink(temp_input.name)
            try:
                os.unlink(temp_output.name)
            except:
                pass
    except Exception as e:
        print(f"❌ Fatal error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

# ============ PDF TO EXCEL ============
@app.route('/convert/pdf-to-excel', methods=['POST'])
def pdf_to_excel():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        print(f"📁 Processing file for Excel: {file.filename}")
        
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_input.name)
        temp_input.close()
        
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        temp_output.close()
        
        try:
            # STEP 1: Coba ekstrak teks biasa
            pdf_reader = PyPDF2.PdfReader(temp_input.name)
            text = ''
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + '\n'
            
            # STEP 2: Kalau kosong, pake OCR
            if not text.strip():
                print("⚠️ No text found for Excel, trying OCR...")
                text = extract_text_with_ocr(temp_input.name)
                if text is None:
                    text = ''
            
            # STEP 3: Buat Excel
            wb = Workbook()
            ws = wb.active
            
            # STEP 4: Deteksi tabel dari teks
            table_data = detect_table_from_text(text)
            
            if table_data:
                print(f"✅ Found {len(table_data)} rows of table data")
                # Tulis ke Excel
                for row_idx, row_data in enumerate(table_data, 1):
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
            else:
                print("⚠️ No table detected, writing as plain text")
                # Fallback: tulis per baris di kolom A
                lines = text.split('\n')
                row_num = 1
                for line in lines:
                    line = line.strip()
                    if line:
                        line = re.sub(r'[^\x00-\x7F]+', ' ', line)
                        ws.cell(row=row_num, column=1, value=line)
                        row_num += 1
            
            wb.save(temp_output.name)
            
            print("✅ Excel file created successfully")
            
            return send_file(
                temp_output.name,
                as_attachment=True,
                download_name='output.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        except Exception as e:
            print(f"❌ Error in conversion: {str(e)}")
            print(traceback.format_exc())
            return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500
        finally:
            os.unlink(temp_input.name)
            try:
                os.unlink(temp_output.name)
            except:
                pass
    except Exception as e:
        print(f"❌ Fatal error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

# ============ WORD TO PDF ============
@app.route('/convert/word-to-pdf', methods=['POST'])
def word_to_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        file.save(temp_input.name)
        temp_input.close()
        
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_output.close()
        
        try:
            doc = Document(temp_input.name)
            doc_pdf = SimpleDocTemplate(temp_output.name, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    story.append(Paragraph(para.text, styles['Normal']))
                    story.append(Spacer(1, 12))
            
            # Tambahkan tabel dari Word
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    data.append(row_data)
                
                if data:
                    t = Table(data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 20))
            
            doc_pdf.build(story)
            
            return send_file(
                temp_output.name,
                as_attachment=True,
                download_name='output.pdf',
                mimetype='application/pdf'
            )
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        finally:
            os.unlink(temp_input.name)
            try:
                os.unlink(temp_output.name)
            except:
                pass
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============ EXCEL TO PDF ============
@app.route('/convert/excel-to-pdf', methods=['POST'])
def excel_to_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        file.save(temp_input.name)
        temp_input.close()
        
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_output.close()
        
        try:
            wb = load_workbook(temp_input.name)
            ws = wb.active
            
            doc_pdf = SimpleDocTemplate(temp_output.name, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            data = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                if any(row_data):
                    data.append(row_data)
            
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
            
            doc_pdf.build(story)
            
            return send_file(
                temp_output.name,
                as_attachment=True,
                download_name='output.pdf',
                mimetype='application/pdf'
            )
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        finally:
            os.unlink(temp_input.name)
            try:
                os.unlink(temp_output.name)
            except:
                pass
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============ HEALTH CHECK ============
@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok', 
        'message': 'Server is running!',
        'ocr_available': OCR_AVAILABLE
    })

@app.route('/', methods=['GET'])
def index():
    return jsonify({
        'name': 'PDF Converter API',
        'version': '1.0',
        'endpoints': [
            '/convert/pdf-to-word',
            '/convert/pdf-to-excel',
            '/convert/word-to-pdf',
            '/convert/excel-to-pdf'
        ],
        'ocr_available': OCR_AVAILABLE
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
